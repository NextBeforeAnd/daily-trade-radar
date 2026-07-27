#!/usr/bin/env python3
"""Build a privacy-scrubbed Word report from validated radar JSON."""

from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


LEVEL_ORDER = {"high": 0, "medium": 1, "low": 2, "watch": 3}
LEVEL_ZH = {"high": "高", "medium": "中", "low": "低", "watch": "观察"}
LEVEL_EN = {"high": "High", "medium": "Medium", "low": "Low", "watch": "Watch"}
STATUS_ZH = {"new": "今日新增", "effective": "已生效", "deadline": "临近截止", "ongoing": "持续关注", "unconfirmed": "待核实"}
STATUS_EN = {"new": "New today", "effective": "Effective", "deadline": "Approaching deadline", "ongoing": "Ongoing", "unconfirmed": "Unconfirmed"}
POLICY_AREA_ZH = {
    "onboarding_kyc": "入驻与身份核验",
    "listing_product_compliance": "刊登与商品合规",
    "pricing_promotions": "定价与促销",
    "fees_commissions": "费用与佣金",
    "fulfillment_logistics": "履约与物流",
    "returns_refunds_aftersales": "退货退款与售后",
    "payments_settlement_tax": "支付、结算与税务",
    "content_ads_affiliate": "内容、广告与联盟",
    "data_privacy_security": "数据、隐私与安全",
    "account_health_enforcement": "账户健康与执法",
    "api_feature_deprecation": "API、功能与停用",
    "other": "其他",
}
LEVEL_FILL = {"high": "FDE8E7", "medium": "FFF1CC", "low": "E8F1FA", "watch": "EEF0F2"}
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 96, 105)
WHITE = RGBColor(255, 255, 255)
TABLE_WIDTH_DXA = 9360
TABLE_WIDTHS_DXA = (900, 1350, 2700, 4410)
RSID_ATTRIBUTE = re.compile(rb"\s+w:rsid(?:R|RDefault|P|RPr|Sect|Del)?=\"[^\"]*\"")
RSID_ELEMENT = re.compile(rb"<w:rsid(?:Root|s)?\b[^>]*(?:/>|>.*?</w:rsids>)")
DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "radar-template.docx"


def load(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("root must be a JSON object")
    return data


def scrub_package(path: Path) -> None:
    """Remove personal/custom metadata and Word revision-session identifiers."""
    descriptor, temporary_name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                if item.filename == "docProps/custom.xml":
                    continue
                payload = source.read(item.filename)
                if item.filename.endswith(".xml"):
                    payload = RSID_ATTRIBUTE.sub(b"", payload)
                    payload = RSID_ELEMENT.sub(b"", payload)
                if item.filename == "[Content_Types].xml":
                    payload = re.sub(rb"<Override[^>]+PartName=\"/docProps/custom.xml\"[^>]*/>", b"", payload)
                if item.filename == "_rels/.rels":
                    payload = re.sub(rb"<Relationship[^>]+Type=\"[^\"]*/custom-properties\"[^>]*/>", b"", payload)
                target.writestr(item, payload)
        temporary.replace(path)
    finally:
        if temporary.exists():
            temporary.unlink()


def set_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: tuple[int, ...]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)


def clear_paragraph(paragraph) -> None:
    """Remove paragraph content while preserving paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def clear_document_body(document: Document) -> None:
    """Remove template body content while preserving section properties and styles."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(document: Document, data: dict, english: bool) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("DAILY TRADE RADAR" if english else "每日外贸雷达")
    set_font(run, size=24, color=DARK_BLUE, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(("Actionable foreign-trade intelligence" if english else "外贸政策与跨境平台行动简报"))
    set_font(run, size=12, color=MUTED)
    metadata = (
        (("Report date", data.get("report_date")), ("Timezone", data.get("timezone")), ("Reporting window starts", data.get("window_start")), ("Search cutoff", data.get("cutoff")), ("Scope", ", ".join(data.get("scope", []))))
        if english else
        (("报告日期", data.get("report_date")), ("时区", data.get("timezone")), ("报告窗口起点", data.get("window_start")), ("检索截止", data.get("cutoff")), ("覆盖范围", "、".join(data.get("scope", []))))
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_font(label_run, size=10.5, bold=True)
        set_font(paragraph.add_run(str(value or "—")), size=10.5)
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(8)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    p_pr.append(borders)


def build_report(data: dict, template_path: Path) -> Document:
    english = str(data.get("language", "")).casefold().startswith("en")
    levels = LEVEL_EN if english else LEVEL_ZH
    statuses = STATUS_EN if english else STATUS_ZH
    events = sorted(data.get("events", []), key=lambda event: (LEVEL_ORDER.get(event.get("level"), 9), -int(event.get("score", 0)), event.get("effective_date") or "9999-99-99"))
    main_events = [event for event in events if event.get("status") != "unconfirmed"]
    watch_events = [event for event in events if event.get("status") == "unconfirmed"]

    if not template_path.is_file():
        raise ValueError(f"template not found: {template_path}")
    document = Document(template_path)
    clear_document_body(document)
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    for paragraph in section.header.paragraphs:
        clear_paragraph(paragraph)
    header = section.header.paragraphs[0]
    for paragraph in section.header.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("Daily Trade Radar" if english else "每日外贸雷达"), size=8.5, color=MUTED)
    for paragraph in section.footer.paragraphs:
        clear_paragraph(paragraph)
    footer = section.footer.paragraphs[0]
    for paragraph in section.footer.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    add_title_block(document, data, english)
    document.add_heading("Today's assessment" if english else "今日判断", level=1)
    if any(event.get("status") == "new" for event in main_events):
        judgment = "Verified new developments were found today. Prioritize the high-risk and urgent actions below." if english else "发现经过核验的今日新增事项，请优先执行下列高风险和紧迫行动。"
    elif main_events:
        judgment = "No verified material new developments were found today. The report retains effective, deadline, and ongoing items." if english else "未发现经过核验的重大今日新增事项；本报告保留已生效、临近截止和持续关注事项。"
    else:
        judgment = "No material items met the main-table evidence standard within this search scope and cutoff." if english else "在本次检索范围和截止时间内，未发现达到主表证据标准的重大事项。"
    document.add_paragraph(judgment)

    document.add_heading("Radar at a glance" if english else "一页雷达", level=1)
    if main_events:
        headers = ("Level", "Status", "Event", "Key impact") if english else ("级别", "状态", "事件", "关键影响")
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        repeat_table_header(table.rows[0])
        for index, header_text in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_shading(cell, "2E74B5")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.add_run(header_text), size=9.5, color=WHITE, bold=True)
        for event in main_events:
            row = table.add_row()
            keep_table_row_together(row)
            values = (levels.get(event.get("level"), str(event.get("level", ""))), statuses.get(event.get("status"), str(event.get("status", ""))), event.get("title", ""), event.get("impact", ""))
            for index, value in enumerate(values):
                cell = row.cells[index]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
                set_font(paragraph.add_run(str(value)), size=9.2, bold=index == 2)
            set_cell_shading(row.cells[0], LEVEL_FILL.get(event.get("level"), "FFFFFF"))
        set_table_geometry(table, TABLE_WIDTHS_DXA)
    else:
        document.add_paragraph("No items met the main-table standard." if english else "本次检索未发现符合主表标准的事项。")

    document.add_heading("Priority actions" if english else "优先行动", level=1)
    actionable = [event for event in main_events if event.get("level") in {"high", "medium"}]
    if actionable:
        for event in actionable:
            items = event.get("action_items") or []
            if items:
                for item in items:
                    paragraph = document.add_paragraph(style="List Number")
                    title_run = paragraph.add_run(f"{event.get('title', '')}: ")
                    set_font(title_run, bold=True)
                    if english:
                        action = str(item.get("action", "")).rstrip("。.;；")
                        evidence = str(item.get("completion_evidence", "")).rstrip("。.;；")
                        paragraph.add_run(
                            f"{item.get('owner')} | {item.get('deadline')}: {action}. "
                            f"Completion evidence: {evidence}."
                        )
                    else:
                        action = str(item.get("action", "")).rstrip("。.;；")
                        evidence = str(item.get("completion_evidence", "")).rstrip("。.;；")
                        paragraph.add_run(
                            f"{item.get('owner')}｜{item.get('deadline')}：{action}；"
                            f"完成凭证：{evidence}。"
                        )
            else:
                paragraph = document.add_paragraph(style="List Number")
                title_run = paragraph.add_run(f"{event.get('title', '')}: ")
                set_font(title_run, bold=True)
                paragraph.add_run(str(event.get("action", "")))
    else:
        document.add_paragraph("No high- or medium-level actions." if english else "暂无高、中等级行动。")

    document.add_heading("Deduplication" if english else "去重说明", level=1)
    matches = data.get("deduplication", {}).get("matches", [])
    removed = sum(item.get("disposition") == "duplicate_removed" for item in matches)
    material = sum(item.get("disposition") in {"material_update", "retained_after_review"} for item in matches)
    operational = sum(item.get("disposition") == "operational_refresh" for item in matches)
    review = sum(item.get("disposition") == "possible_update" for item in matches)
    if english:
        document.add_paragraph(
            f"Compared with the previous report, {removed} unchanged duplicate(s) were removed; "
            f"{material} material update(s) and {operational} operational refresh(es) were retained; "
            f"{review} item(s) still require review."
        )
    else:
        document.add_paragraph(
            f"与上一期比较后移除 {removed} 条无实质变化的重复事项；"
            f"保留 {material} 条实质更新事项和 {operational} 条运营节点刷新事项；"
            f"另有 {review} 条仍待复核。"
        )

    document.add_heading("Event details and official sources" if english else "事件详情与官方来源", level=1)
    for index, event in enumerate(main_events, 1):
        document.add_heading(f"S{index} | {event.get('title', '')}", level=2)
        detail_rows = (
            (("Jurisdiction", event.get("jurisdiction")), ("Authority", event.get("authority")), ("Published", event.get("published_date")), ("Published at", event.get("published_at")), ("Effective", event.get("effective_date")), ("Effective at", event.get("effective_at")), ("Deadline", event.get("deadline")), ("Deadline at", event.get("deadline_at")), ("Source timezone", event.get("source_timezone")), ("Summary", event.get("summary")), ("Action", event.get("action")))
            if english else
            (("司法辖区", event.get("jurisdiction")), ("主管机构", event.get("authority")), ("发布日期", event.get("published_date")), ("发布时间", event.get("published_at")), ("生效日期", event.get("effective_date")), ("生效时间", event.get("effective_at")), ("截止日期", event.get("deadline")), ("截止时间", event.get("deadline_at")), ("来源时区", event.get("source_timezone")), ("摘要", event.get("summary")), ("行动", event.get("action")))
        )
        policy = event.get("platform_policy")
        if isinstance(policy, dict):
            area = policy.get("policy_area") if english else POLICY_AREA_ZH.get(policy.get("policy_area"), policy.get("policy_area"))
            platform_rows = (
                (("Platform / market", f"{policy.get('platform')} / {policy.get('seller_market')}"), ("Program", policy.get("program")), ("Policy area", area), ("Change type", policy.get("change_type")), ("Seller scope", policy.get("seller_scope")), ("Previous state", policy.get("previous_state")), ("New state", policy.get("new_state")), ("Enforcement consequence", policy.get("enforcement_consequence")), ("Backend verification required", "Yes" if policy.get("backend_verification_required") else "No"))
                if english else
                (("平台 / 市场", f"{policy.get('platform')} / {policy.get('seller_market')}"), ("项目 / 模式", policy.get("program")), ("政策领域", area), ("变化类型", policy.get("change_type")), ("卖家范围", policy.get("seller_scope")), ("原状态", policy.get("previous_state")), ("新状态", policy.get("new_state")), ("执法后果", policy.get("enforcement_consequence")), ("需后台核验", "是" if policy.get("backend_verification_required") else "否"))
            )
            detail_rows = tuple(platform_rows) + tuple(detail_rows)
        for label, value in detail_rows:
            if value in (None, ""):
                continue
            paragraph = document.add_paragraph()
            set_font(paragraph.add_run(f"{label}: "), bold=True)
            paragraph.add_run(str(value))
        source_paragraph = document.add_paragraph()
        set_font(source_paragraph.add_run(("Official source: " if english else "官方来源：")), bold=True)
        add_hyperlink(source_paragraph, str(event.get("source_title") or event.get("source_url") or "Source"), str(event.get("source_url") or ""))
        source_paragraph.add_run((f", retrieved {event.get('retrieved_date')}." if english else f"，检索于 {event.get('retrieved_date')}。"))

    document.add_heading("Unconfirmed watchlist" if english else "待核实观察", level=1)
    if watch_events:
        for event in watch_events:
            paragraph = document.add_paragraph(style="List Bullet")
            set_font(paragraph.add_run(f"{event.get('title', '')}: "), bold=True)
            paragraph.add_run(str(event.get("summary", "")))
    else:
        document.add_paragraph("None." if english else "无。")

    document.add_heading("Coverage gaps" if english else "覆盖缺口", level=1)
    gaps = data.get("coverage_gaps", [])
    if gaps:
        for gap in gaps:
            document.add_paragraph(str(gap), style="List Bullet")
    else:
        document.add_paragraph("No known coverage gaps." if english else "无已知覆盖缺口。")

    coverage_ledger = data.get("coverage_ledger", [])
    if coverage_ledger:
        document.add_heading("Platform coverage ledger" if english else "平台覆盖台账", level=1)
        for entry in coverage_ledger:
            heading = f"{entry.get('platform')} / {entry.get('seller_market')} — {entry.get('program')}"
            document.add_heading(heading, level=2)
            checks = (
                f"Public update: {'Yes' if entry.get('public_update_checked') else 'No'}; "
                f"current policy: {'Yes' if entry.get('current_policy_checked') else 'No'}; "
                f"dashboard: {'Yes' if entry.get('dashboard_checked') else 'No'}."
                if english else
                f"公告核验：{'是' if entry.get('public_update_checked') else '否'}；"
                f"现行政策核验：{'是' if entry.get('current_policy_checked') else '否'}；"
                f"后台核验：{'是' if entry.get('dashboard_checked') else '否'}。"
            )
            document.add_paragraph(checks)
            result_label = "Access result" if english else "访问结果"
            checked_label = "Checked at" if english else "核验时间"
            access_result = entry.get("access_result")
            if not english:
                access_result = {
                    "opened": "已打开",
                    "login_required": "需要登录",
                    "blocked": "访问受阻",
                    "not_checked": "未核验",
                }.get(access_result, access_result)
            document.add_paragraph(f"{result_label}: {access_result} | {checked_label}: {entry.get('checked_at')}")
            lookback_label = "Platform lookback start" if english else "平台回看起点"
            document.add_paragraph(f"{lookback_label}: {entry.get('lookback_start')}")
            sources_label = "Sources opened" if english else "已打开来源"
            document.add_paragraph(sources_label)
            for source in entry.get("sources_checked", []):
                snapshot = source.get("snapshot") if isinstance(source.get("snapshot"), dict) else None
                snapshot_text = ""
                if snapshot:
                    snapshot_label = "snapshot" if english else "快照"
                    snapshot_text = (
                        f" · {snapshot_label}: {snapshot.get('change_status')}"
                        f" · {snapshot.get('diff_summary')}"
                    )
                document.add_paragraph(
                    f"{source.get('source_type')} / {source.get('result')} · "
                    f"{source.get('checked_at')} · {source.get('notes')} · {source.get('url')}"
                    f"{snapshot_text}",
                    style="List Bullet",
                )
            for gap in entry.get("gaps", []):
                document.add_paragraph(str(gap), style="List Bullet")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.keep_together = True
    set_font(note.add_run("Important: " if english else "重要提示："), size=9, color=MUTED, bold=True)
    set_font(note.add_run("This report is an operational research aid, not legal, tax, customs, or sanctions advice." if english else "本报告用于业务研究辅助，不构成法律、税务、海关或制裁专业意见。"), size=9, color=MUTED)
    return document


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    try:
        data = load(args.input)
        document = build_report(data, args.template)
        document.core_properties.title = "Daily Trade Radar" if str(data.get("language", "")).casefold().startswith("en") else "每日外贸雷达"
        document.core_properties.subject = "Actionable foreign-trade intelligence"
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.core_properties.keywords = "trade, customs, compliance, marketplace"
        args.output.parent.mkdir(parents=True, exist_ok=True)
        document.save(args.output)
        scrub_package(args.output)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}")
        return 2
    print(f"WROTE: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
